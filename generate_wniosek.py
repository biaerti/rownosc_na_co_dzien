from docx import Document
from docx.shared import Pt
import os

folder = r'c:\VisualStudioFolders\rownosc-na-co-dzien'
subfolder_name = [x for x in os.listdir(folder) if 'atno' in x][0]
output_path = os.path.join(folder, subfolder_name, 'wniosek_grudzien25_luty26_DRAFT.docx')

doc = Document()

def heading(text, level=1):
    return doc.add_heading(text, level=level)

def body(text):
    p = doc.add_paragraph(text)
    return p

# Title
doc.add_heading('Wniosek o platnosc do projektu', 0)
doc.add_heading('FEDS.07.03-IP.02-0039/25', 1)
doc.add_heading('Rownosc na co dzien - przeciwdzialanie dyskryminacji i przemocy', 1)
body('Wniosek za okres od 2025-12-01 do 2026-02-28')
body('Rodzaj wniosku: Rozliczajacy zaliczke, Sprawozdawczy')
body('Beneficjent: Fundacja Pretium')
doc.add_paragraph()

heading('POSTEP RZECZOWY', 1)

# ------ ZADANIE 1 ------
heading('Zadanie 1 Kampania informacyjna na rzecz zapobiegania dyskryminacji i przemocy', 2)
heading('Stan realizacji:', 3)
z1 = (
    "W ramach realizacji Zadania 1 kontynuowano dzialania kampanii informacyjnej na platformach "
    "spolecznosciowych oraz stronie internetowej projektu. W okresie sprawozdawczym publikowane byly tresci "
    "edukacyjne na platformach Facebook, Instagram oraz TikTok, dotyczace tematyki rownosci szans, "
    "przeciwdzialania dyskryminacji i przemocy \u2013 w tym materialy promujace system RESQL jako narzedzie "
    "anonimowego zglaszania przemocy i dyskryminacji w srodowisku szkolnym oraz materialy promujace "
    "dzialalnosc Centrum Wsparcia projektu.\n\n"
    "Tresci publikowane w ramach kampanii mialy charakter edukacyjny i informacyjny, dostosowany do szerokiego "
    "grona odbiorcow. Kampania prowadzona byla zgodnie z zasadami rownosci szans i niedyskryminacji.\n\n"
    "W ramach dzialania Publikacja tresci edukacyjnych na platformach spolecznosciowych poniesiono wydatki "
    "zgodnie z harmonogramem \u2013 faktury nr FNP4 1/12/25 (grudzien 2025) oraz FNP4 2/01/26 (styczen 2026) "
    "na laczna kwote 3 900,00 zl.\n\n"
    "Trwaja prace nad przygotowaniem kolejnych filmow edukacyjnych (Film 2 i Film 3 z zaplanowanych 4). "
    "Ponadto Fundacja nawiazala kontakt z ekspertami w celu przygotowania podcastow z udzialem specjalistow "
    "(dzialanie 1.2, 2 szt. x 5 000 zl) \u2013 dzialanie planowane do realizacji w kolejnych miesiacach.\n\n"
    "Realizacja Zadania 1 przebiega zgodnie z harmonogramem. Kampania przyczynia sie do wzrostu swiadomosci "
    "spolecznej w zakresie rownosci szans, dostepnosci wsparcia dla osob doswiadczajacych dyskryminacji "
    "oraz zasiegu dzialan projektowych."
)
body(z1)
doc.add_paragraph()

# ------ ZADANIE 2 ------
heading('Zadanie 2 Konferencje', 2)
heading('Stan realizacji:', 3)
z2 = (
    "W okresie sprawozdawczym, w dniach 15\u201316 stycznia 2026 roku, zorganizowano dwudniowa konferencje "
    "pt. \u201eRownosc w ochronie zdrowia\u201d skierowana do kadry organizacyjnej, administracyjnej i zarzadzajacej "
    "placowek medycznych z wojewodztwa dolnoslaskiego. Wydarzenie odbylo sie w Hotelu Dziki Potok w Karpaczu.\n\n"
    "Konferencja zgromadzila 44 uczestnikow (42 kobiety i 2 mezczyzn) \u2013 przedstawicieli kadry zarzadzajacej "
    "i administracyjnej systemu ochrony zdrowia. Rekrutacja uczestnikow prowadzona byla zgodnie z zasada "
    "rownosci szans, z priorytetem dla osob z obszarow wiejskich oraz osob z orzeczeniem o niepelnosprawnosci. "
    "Fundacja Pretium, realizujac rowniez szkolenia medyczne jako jedna z glownych dzialalnosci statutowych, "
    "skutecznie dotarla do placowek medycznych z obszaru dolnoslaskiego i zapewniala udzial wykwalifikowanej "
    "kadry.\n\n"
    "Program konferencji obejmowal m.in.:\n"
    "\u2022 Sesje inauguracyjna: wprowadzenie do tematyki rownosci w ochronie zdrowia (dr n. o zdr. Monika Tomaszewska),\n"
    "\u2022 Wyklad ekspercki: \u201eDyskryminacja w systemie ochrony zdrowia \u2013 identyfikacja problemow i wyzwan\u201d (mgr Anna Daria Nowicka),\n"
    "\u2022 Panel dyskusyjny: \u201eWdrazanie polityki rownosci w placowkach medycznych\u201d,\n"
    "\u2022 Studium przypadku: analiza sytuacji konfliktowych zwiazanych z nierownym traktowaniem,\n"
    "\u2022 Warsztaty interaktywne: \u201ePraktyczne narzedzia przeciwdzialania dyskryminacji\u201d (mgr Iwona Wojnowska),\n"
    "\u2022 Wypracowanie rekomendacji i planow dzialan na rzecz rownosci w ochronie zdrowia,\n"
    "\u2022 Uroczysta kolacja integracyjna i sesje networkingowe.\n\n"
    "Wszyscy uczestnicy wypelnili pre-test i post-test mierzacy zmiane poziomu wiedzy w zakresie rownosci szans "
    "i niedyskryminacji. Konferencja zostala zrealizowana zgodnie z harmonogramem i zalozeniami wniosku "
    "o dofinansowanie, z poszanowaniem zasad rownosci szans i dostepnosci.\n\n"
    "Realizacja konferencji wyczerpuje zaplanowane dzialanie w ramach Zadania 2 (konferencja dla kadry "
    "systemu ochrony zdrowia). Laczny koszt organizacji wyniose 54 900,00 zl "
    "(sala: 14 000 zl, prelegenci: 16 000 zl, obiad: 7 800 zl, kolacja: 4 200 zl, "
    "nocleg ze sniadaniem: 12 900 zl), na podstawie faktury FNP4 1/01/26 (Fundacja Diligo)."
)
body(z2)
doc.add_paragraph()

# ------ ZADANIE 3 ------
heading('Zadanie 3 Wdrozenie systemu RESQL w szkolach', 2)
heading('Stan realizacji:', 3)
z3 = (
    "W okresie sprawozdawczym kontynuowano intensywne dzialania zmierzajace do wdrozenia systemu RESQL "
    "w 30 szkolach na terenie wojewodztwa dolnoslaskiego. Realizacja Zadania 3 przebiega zgodnie "
    "z harmonogramem i we wspolpracy z partnerem \u2013 RESQL Sp. z o.o.\n\n"
    "W dniu 27 stycznia 2026 r. przeprowadzono spotkania online (Google Meet) z pierwsza tura szkol "
    "zakwalifikowanych do projektu, poswiecone procesowi uzupelniania dokumentacji wdrozeniowej. "
    "Spotkaniami objeto nastepujace szkoly:\n"
    "\u2022 Centrum Ksztalcenia Zawodowego i Ustawicznego w Strzelinie,\n"
    "\u2022 Zespol Szkolno-Przedszkolny w Iwinach,\n"
    "\u2022 Zespol Szkolno-Przedszkolny w Pustkowie Zurawskim,\n"
    "\u2022 Zespol Szkolno-Przedszkolny w Bielanach Wroclawskich,\n"
    "\u2022 Szkola Podstawowa im. Bohaterow Westerplatte w Chobieni,\n"
    "\u2022 Szkola Podstawowa Nr 78,\n"
    "\u2022 Szkola Podstawowa im. B. Chrobrego w Zorawinie,\n"
    "\u2022 Szkola Podstawowa im. Boleslawa Chrobrego w Zmigrodzie.\n\n"
    "W dniu 26 lutego 2026 r. odbyly sie analogiczne spotkania online z druga tura szkol:\n"
    "\u2022 Liceum Ogolnoksztalcace Nr XV,\n"
    "\u2022 Liceum Ogolnoksztalcace nr XIV im. Polonii Belgijskiej,\n"
    "\u2022 Zespol Szkolno-Przedszkolny nr 12,\n"
    "\u2022 Zespol Szkol im. Jedrzeja Sniadeckiego w Zarowie,\n"
    "\u2022 Powiatowy Zespol Szkol nr 1 w Krzyzowicach,\n"
    "\u2022 Szkola Podstawowa z Oddzialami Integracyjnymi w Lomnicy,\n"
    "\u2022 Zespol Szkolno-Przedszkolny nr 1.\n\n"
    "W dniu 3 lutego 2026 r. przeprowadzono dodatkowe spotkanie z pierwsza tura szkol bedacych na etapie "
    "finalizowania dokumentacji, w celu omowienia szczegolnych aspektow procesu wdrozeniowego.\n\n"
    "W celu ulatwienia plakowkom prawidlowego wypelnienia dokumentacji wdrozeniowej udostepniono dedykowany "
    "folder na Google Drive (https://drive.google.com/drive/folders/1elDURMJ1RsJ4cQ26PDk0AYAvWZC_Xc_h) "
    "zawierajacy wszystkie niezbedne dokumenty oraz film instruktazowy prezentujacy krok po kroku "
    "sposob ich uzupelnienia.\n\n"
    "Na dzien sprawozdawczy zebrano kompleta dokumentacje wdrozeniowa od 15 szkol \u2013 co stanowi "
    "polowe zaplanowanego wskaznika. Pierwsze szkoly zakonczyly etap zbierania dokumentow i przygotowuja "
    "sie do realizacji szkolen i wdroz systemu RESQL. Lada chwila partner RESQL przystapi do "
    "przeprowadzenia pierwszej serii szkolen wdrozeniowych w zakwalifikowanych placowkach.\n\n"
    "Fundacja Pretium prowadzila rownolegla dzialania wspierajace rekrutacje i kontakt ze szkolami, "
    "w tym kontynuowala relacje nawiazane podczas konferencji dla kadry oswiaty (Zadanie 2)."
)
body(z3)
doc.add_paragraph()

# ------ ZADANIE 4 ------
heading('Zadanie 4 Centrum wsparcia \u2013 wsparcie psychologiczne i pedagogiczne dla osob dyskryminowanych', 2)
heading('Stan realizacji:', 3)
z4 = (
    "W okresie sprawozdawczym kontynuowano funkcjonowanie Centrum Wsparcia przy ul. Bierutowskiej 57\u201359 "
    "we Wroclawiu, w ramach ktorego swiadczone jest kompleksowe wsparcie dla osob doswiadczajacych "
    "dyskryminacji i przemocy.\n\n"
    "Centrum Wsparcia dzialalo zgodnie z zalozeniami projektu. Zatrudnieni pedagodzy (2 osoby na pelen etat) "
    "realizowali biezace wsparcie uczestnikow projektu, obslugiwali telefon zaufania (nr 690 259 392, "
    "czynny pn.\u2013pt. w godz. 7:00\u201315:30) oraz prowadzili pierwsze kontakty i rozmowy pomocowe. "
    "Psycholog zatrudniony na podstawie umowy cywilnoprawnej prowadzil indywidualne konsultacje "
    "psychologiczne oraz interwencje kryzysowe w wymiarze 4 godzin miesiecznie, zgodnie z umowa "
    "(14h/miesiac na caly projekt). Wsparcie psychologiczne realizowane jest przez specjaliste posiadajacego "
    "wymagane wyksztalcenie kierunkowe oraz ukonczony kurs wsparcia kryzysowego.\n\n"
    "W ramach Centrum Wsparcia nawiazano wspolprace z:\n"
    "\u2022 Portalem Pomocy Psychologicznej psychotekst.pl (Wroclaw) \u2013 portal umie"
    "scil materialy promocyjne Centrum Wsparcia na swojej stronie internetowej oraz na Facebooku,\n"
    "\u2022 Schroniskiem Swietego Brata Alberta dla Bezdomnych Kobiet we Wroclawiu \u2013 "
    "schronisko udostepnilo plakaty Centrum Wsparcia oraz opublikowalo informacje na swojej stronie internetowej.\n\n"
    "Podjete dzialania promocyjne przyczyniaja sie do zwiekszenia rozpoznawalnosci i dostepnosci Centrum "
    "Wsparcia wsrod grup szczegolnie narazonych na dyskryminacje i przemoc.\n\n"
    "W ramach Centrum Wsparcia realizowano rowniez warsztaty dla kobiet z zakresu budowania pewnosci siebie "
    "oraz negocjacji (wynagrodzenia, warunki pracy). W okresie sprawozdawczym wsparciem w ramach warsztatow "
    "objetych zostalo lacznie 15 uczestniczek. Zajecia prowadzone byly w malych grupach (od 2 do 5 osob), "
    "co umozliwilo indywidualne podejscie do potrzeb kazdej uczestniczki i objetych zostalo caly zakres "
    "tematyczny (pewnosc siebie, negocjacje). Warsztaty realizowane sa przez pedagoga z wyksztalceniem "
    "psychologicznym, z uwzglednieniem praktycznego wymiaru wsparcia \u2013 w tym m.in. pomocy "
    "w przygotowaniu dokumentow aplikacyjnych z wykorzystaniem nowoczesnych narzedzi cyfrowych, "
    "gdyz kobiety sa szczegolnie narazone na dyskryminacje na rynku pracy. Warsztaty ciesza sie "
    "bardzo pozytywnymi opiniami uczestniczek.\n\n"
    "Wszystkie dzialania Centrum Wsparcia realizowane sa zgodnie z zasadami rownosci szans "
    "i niedyskryminacji oraz z poszanowaniem prywatnosci i godnosci uczestniczek i uczestnikow."
)
body(z4)
doc.add_paragraph()

# ------ ZADANIE 5 ------
heading('Zadanie 5 Szkolenia zwiekszajace kompetencje pracownikow sluzby zdrowia', 2)
heading('Stan realizacji:', 3)
body('Szkolenie nie bylo realizowane w okresie sprawozdawczym.')
doc.add_paragraph()

# ------ KOSZTY POSREDNIE ------
heading('Zadanie Koszty posrednie', 2)
heading('Stan realizacji:', 3)
kp = (
    "Wydatki w ramach kosztow posrednich w biezacym okresie sprawozdawczym byly ponoszone zgodnie "
    "z aktualnie obowiazujacymi regulacjami prawnymi oraz wytycznymi w zakresie kwalifikowalnosci "
    "kosztow w perspektywie finansowej UE na lata 2021\u20132027.\n\n"
    "Personel realizujacy projekt kontynuowal stosowanie zasad rownosci szans kobiet i mezczyzn "
    "oraz dostepnosci dzialan projektowych dla osob z niepelnosprawnosciami. Czlonkowie zespolu "
    "wykorzystywali w codziennej pracy dokument pn. \u201eWytyczne w zakresie realizacji zasady "
    "rownosci szans i niedyskryminacji (...)\u201d, traktujac go jako praktyczne narzedzie wspierajace "
    "planowanie i realizacje dzialan.\n\n"
    "W okresie sprawozdawczym w biurze projektu odbywaly sie systematyczne spotkania zespolu, "
    "podczas ktorych:\n"
    "\u2022 omawiano biezacy postep w realizacji harmonogramu,\n"
    "\u2022 analizowano stopien osiagniecia wskaznikow produktu i rezultatu,\n"
    "\u2022 identyfikowano trudnosci oraz wypracowywano rozwiazania dostosowane do zmieniajacych sie "
    "potrzeb uczestnikow projektu.\n\n"
    "W dalszym ciagu stosowano zasady rownosciowego zarzadzania projektem. Czlonkom zespolu "
    "umozliwiono prace w elastycznych godzinach oraz czesciowo w trybie zdalnym, co pozwolilo na "
    "lepsze godzenie obowiazkow zawodowych i prywatnych.\n\n"
    "W ramach kosztow posrednich finansowano w szczegolnosci:\n"
    "\u2022 koszty zwiazane z funkcjonowaniem i biezacym utrzymaniem biura projektu,\n"
    "\u2022 obsluge administracyjno-finansowa projektu,\n"
    "\u2022 zakup materialow biurowych i eksploatacyjnych niezbednych do sprawnej realizacji zadan.\n\n"
    "Dzialania na rzecz zrownowazonego rozwoju:\n"
    "Zespol projektowy konsekwentnie dbaal o ograniczenie negatywnego wplywu na srodowisko, w tym m.in.:\n"
    "\u2022 ograniczano druk materialow wylacznie do sytuacji niezbednych, stosujac druk dwustronny i tryb czarno-bialy,\n"
    "\u2022 prowadzono segregacje odpadow w biurze (ze szczegolnym uwzglednieniem makulatury),\n"
    "\u2022 racjonalnie korzystano z energii elektrycznej i urzadzen biurowych,\n"
    "\u2022 preferowano naturalne oswietlenie i wentylacje w miejscu pracy, ograniczajac uzycie klimatyzacji.\n\n"
    "Beneficjent oswiadcza, ze konta dedykowane do projektow zarowno Lidera jak i partnerow "
    "sa kontami nieoprocentowanymi."
)
body(kp)
doc.add_paragraph()

# ------ PROBLEMY ------
heading('Problemy napotkane w trakcie realizacji projektu', 2)
body('W okresie sprawozdawczym nie napotkano problemow.')
doc.add_paragraph()

# ------ PLANOWANY PRZEBIEG ------
heading('Planowany przebieg realizacji projektu', 2)
pp = (
    "W kolejnych miesiacach projekt bedzie realizowany zgodnie z harmonogramem, z naciskiem na "
    "nastepujace dzialania:\n\n"
    "Centrum Wsparcia (Zadanie 4): planowane jest dalsze intensywne rozwijanie dzialalnosci Centrum "
    "Wsparcia, w tym zwiekszenie liczby indywidualnych konsultacji psychologicznych i sesji interwencji "
    "kryzysowej. Po zakonczeniu intensywnej fazy operacyjnej (konferencje, wdrozenia RESQL) zespol "
    "projektowy zwiekszy zaangazowanie w promocje i realizacje uslug wsparcia bezposredniego. "
    "Centrum Wsparcia bedzie aktywnie promowane w ramach kampanii informacyjnej, w tym poprzez "
    "dzialania na platformach spolecznosciowych (Facebook, Instagram, TikTok). Planowane jest rowniez "
    "uruchomienie grupy wsparcia \u2013 cyklicznych spotkan dla osob doswiadczajacych dyskryminacji, "
    "promowanej m.in. na Facebooku projektu.\n\n"
    "Kampania informacyjna (Zadanie 1): planowane jest wejscie na platformy spolecznosciowe kolejnych "
    "2 filmow edukacyjnych (Film 2 i Film 3 z zaplanowanych 4). W ramach dzialania 1.2 trwa "
    "nawiazywanie kontaktow z ekspertami i specjalistami z obszaru rownosci i niedyskryminacji "
    "w celu przygotowania 2 podcastow z udzialem ekspertow (2 szt. x 5 000 zl). Planowana jest "
    "kontynuacja i intensyfikacja kampanii sponsorowanej na platformach FB, IG i TikTok, z naciskiem "
    "na promocje Centrum Wsparcia i jego uslug.\n\n"
    "Wdrozenie RESQL w szkolach (Zadanie 3): planowane jest domkniecie procesu zbierania dokumentacji "
    "od pozostalych szkol oraz przeprowadzenie przez partnera RESQL Sp. z o.o. serii szkolen "
    "wdrozeniowych \u2013 dla dyrekcji, interwentow, kadry pedagogicznej oraz uczniow \u2013 "
    "we wszystkich 30 zakwalifikowanych placowkach.\n\n"
    "Szkolenia dla pracownikow sluzby zdrowia (Zadanie 5): planowane jest przygotowanie i realizacja "
    "10 edycji szkolen jednodniowych dla pracownikow sluzby zdrowia, zgodnie z harmonogramem projektu.\n\n"
    "Kolejne warsztaty dla kobiet beda kontynuowane w ramach Centrum Wsparcia."
)
body(pp)
doc.add_paragraph()

# ------ WSKAZNIKI ------
heading('WSKAZNIKI \u2013 wartosci do uzupelnienia w systemie CST', 1)
body(
    "Uwaga: wartosci w tym okresie (nowe w biezacym sprawozdaniu) / narastajaco od poczatku projektu.\n\n"
    "1. Osoby obj\u0119te wsparciem (rownosc K/M):\n"
    "   Ogolnie: 63 / narastajaco 132 | K: 57 / nar. 119 | M: 6 / nar. 13\n\n"
    "2. Osoby ponizej 18 lat:\n"
    "   Ogolnie: 0 | K: 0 | M: 0\n\n"
    "3. Osoby 18\u201329 lat:\n"
    "   Ogolnie: 8 / nar. 13 | K: 6 / nar. 11 | M: 2 / nar. 2\n\n"
    "4. Osoby 55+:\n"
    "   Ogolnie: 10 / nar. 26 | K: 8 / nar. 21 | M: 2 / nar. 5\n\n"
    "5. Podmioty publiczne:\n"
    "   Ogolnie: 15 / nar. 15 (szkoly, ktore zlozily dokumenty wdrozeniowe)\n\n"
    "6. Obiekty dostosowane do potrzeb ON: 0\n\n"
    "7. Projekty z racjonalnymi usprawnieniami: 0\n\n"
    "8. Osoby z obszarow wiejskich:\n"
    "   Ogolnie: 5 / nar. 29 | K: 5 / nar. 27 | M: 0 / nar. 2\n\n"
    "9. Osoby obj\u0119te wsparciem rownosc szans/niedyskryminacja:\n"
    "   Ogolnie: 63 / nar. 132 | K: 57 / nar. 119 | M: 6 / nar. 13\n\n"
    "10. Mniejszosci (w tym Romowie): 0\n\n"
    "11. Osoby z niepelnosprawnosciami:\n"
    "    Ogolnie: 1 / nar. 1 | K: 1 | M: 0\n\n"
    "12. Osoby obcego pochodzenia:\n"
    "    Ogolnie: 4 / nar. 4 | K: 4 | M: 0\n"
    "    [zweryfikuj plec na podstawie danych w SL \u2013 CSV pokazuje 4 osoby 'obcego pochodzenia']\n\n"
    "13. Osoby w kryzysie bezdomnosci: 0\n\n"
    "14. Osoby z krajow trzecich: 0"
)
doc.add_paragraph()

# ------ ZESTAWIENIE WYDATKOW ------
heading('ZESTAWIENIE WYDATKOW W OKRESIE SPRAWOZDAWCZYM', 1)
body(
    "Ponizej wydatki wykazane w tym okresie (grudzien 2025 \u2013 styczen 2026):\n\n"
    "ZADANIE 1 \u2013 Kampania informacyjna:\n"
    "\u2022 FNP4 1/12/25 \u2013 Publikacja tresci edukacyjnych (grudzien 2025): 1 950,00 zl | PandA spolka cywilna\n"
    "\u2022 FNP4 2/01/26 \u2013 Publikacja tresci edukacyjnych (styczen 2026): 1 950,00 zl | Patryk Krzyzanowski\n"
    "Suma Zadanie 1: 3 900,00 zl\n\n"
    "ZADANIE 2 \u2013 Konferencja dla kadry systemu ochrony zdrowia (Karpacz, 15\u201316.01.2026):\n"
    "\u2022 FNP4 1/01/26 (Fundacja Diligo, 0001/01/26/FVS) \u2013 Organizacja konferencji:\n"
    "  \u2013 sala: 14 000,00 zl\n"
    "  \u2013 prelegenci: 16 000,00 zl\n"
    "  \u2013 obiad: 7 800,00 zl\n"
    "  \u2013 kolacja: 4 200,00 zl\n"
    "  \u2013 nocleg ze sniadaniem: 12 900,00 zl\n"
    "Suma Zadanie 2: 54 900,00 zl\n\n"
    "ZADANIE 4 \u2013 Centrum Wsparcia:\n"
    "\u2022 FNP4 3/01/26 (Fundacja Diligo, 0002/01/26/FVS) \u2013 Wsparcie psychologiczne \u2013 interwent kryzysowy: 880,00 zl\n"
    "\u2022 LPP4 1/12/25 \u2013 Wynagrodzenia \u2013 umowa o prace za XII (platne I 2026): 12 975,79 zl\n"
    "\u2022 LPP4 1/01/26 \u2013 Wynagrodzenia \u2013 umowa o prace za I (platne II 2026): 12 451,00 zl\n"
    "Suma Zadanie 4: 26 306,79 zl\n\n"
    "LACZNIE WYDATKI BEZPOSREDNIE: 85 106,79 zl"
)

doc.save(output_path)
print('Zapisano:', output_path)
